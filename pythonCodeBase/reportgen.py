'''
Created on Nov 9, 2023

@author: Chathinthaka
'''
import sys
import os
#import re
from collections import OrderedDict
from docutils.nodes import row
#from future.backports.test.pystone import FALSE
sys.path.append(os.path.join(os.path.dirname(os.getcwd()), '360Survey/Db'))
print(sys.path)
from MySQLConnection import MySQLConnection
import getopt
import mysql.connector
from enum import IntEnum, Enum
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.patches import Circle, RegularPolygon
from matplotlib.path import Path
from matplotlib.projections.polar import PolarAxes
from matplotlib.projections import register_projection
from matplotlib.spines import Spine
from matplotlib.transforms import Affine2D
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.dml.color import ColorFormat
from docx.shared import Cm, Inches
from docx.shared import RGBColor
from docx.oxml.table import CT_Row, CT_Tc
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


# db_host= 'localhost'
# db_name= '360_survey_schema'
# db_user = 'root'
# db_password = '1qaz2wsx!@'

db_host= '50.87.232.129'
db_name= 'thrwcons_hayleys_360'
db_user = 'thrwcons_root'
db_password = '1rkLhgDMoY2n'


DisplayGraphsFlag = False
KeepTempsFlag = False

class SurveyState(IntEnum):
  INVALID = 0
  ASSIGN = 1
  INPROGRESS = 2
  COMPLETE =3

  
SurveyState2String = {
  SurveyState.INVALID : "InvalidState",
  SurveyState.ASSIGN : "Assign",
  SurveyState.INPROGRESS : "InProgress",
  SurveyState.COMPLETE : "Complete"
  }

#Need to initialize this from the config file or QuestionList Json file itself
RatingLevels = OrderedDict([(1 , 'Not demonstrated it at all'),
    (2 , '3 out of 10 times the individual has demonstrated it'),
    (3 , '4-7 times out of 10 the individual has demonstrated it'),
    (4 , '8-9 times out of 10 the individual has demonstrated it'),
    (5 , 'Individual has demonstrated in every interaction')])


RatingLevel2Label = OrderedDict([(1 , "Not demonstrated it\nat all"),
    (2 , "3 out of 10 times the\nindividual has demonstrated it"),
    (3 , "4-7 times out of 10 the\nindividual has demonstrated it"),
    (4 , "8-9 times out of 10 the\nindividual has demonstrated it"),
    (5 , "Individual has demonstrated\nin every interaction")])

  
class SurveyorType(IntEnum):
  SELF = 0
  REPORTING_MANAGER = 1
  DIRECT_REPORT = 2
  PEER = 3
  INTERNAL_SURVEYOR = 4
  EXTERNAL_SURVEYOR = 5


class SurveyorTypeKey(str, Enum):
  SELF = "Self"
  REPORTING_MANAGER = "ReportingManager"
  DIRECT_REPORT = "DirectReport"
  PEER = "Peer"
  INTERNAL_SURVEYOR = "InternalSurveyor"
  EXTERNAL_SURVEYOR = "ExternalSurveyor"


SurveyorType2SurveyorTypeKey = {
  SurveyorType.SELF : SurveyorTypeKey.SELF,
  SurveyorType.REPORTING_MANAGER : SurveyorTypeKey.REPORTING_MANAGER,
  SurveyorType.DIRECT_REPORT : SurveyorTypeKey.DIRECT_REPORT,
  SurveyorType.PEER : SurveyorTypeKey.PEER,
  SurveyorType.INTERNAL_SURVEYOR : SurveyorTypeKey.INTERNAL_SURVEYOR,
  SurveyorType.EXTERNAL_SURVEYOR : SurveyorTypeKey.EXTERNAL_SURVEYOR  
}

  
SurveyorTypeKey2SurveyorType = {
  SurveyorTypeKey.SELF : SurveyorType.SELF,
  SurveyorTypeKey.REPORTING_MANAGER : SurveyorType.REPORTING_MANAGER, 
  SurveyorTypeKey.DIRECT_REPORT : SurveyorType.DIRECT_REPORT,
  SurveyorTypeKey.PEER : SurveyorType.PEER,
  SurveyorTypeKey.INTERNAL_SURVEYOR : SurveyorType.INTERNAL_SURVEYOR, 
  SurveyorTypeKey.EXTERNAL_SURVEYOR : SurveyorType.EXTERNAL_SURVEYOR
}


SurveyorTypeKey2Label = {
  SurveyorTypeKey.SELF :  'Self Ratings',
  SurveyorTypeKey.REPORTING_MANAGER : 'Reporting Manager Avg',
  SurveyorTypeKey.DIRECT_REPORT : 'Direct Reports Avg',
  SurveyorTypeKey.PEER : 'Peer Avg',
  SurveyorTypeKey.INTERNAL_SURVEYOR : 'Internal Surveyor Avg',
  SurveyorTypeKey.EXTERNAL_SURVEYOR : 'External Surveyor Avg'  
}
  
    
ColorMap = {
  SurveyorType.SELF : 'k',
  SurveyorType.REPORTING_MANAGER : 'r',
  SurveyorType.DIRECT_REPORT : 'b',
  SurveyorType.PEER : 'g',
  SurveyorType.INTERNAL_SURVEYOR : 'c',
  SurveyorType.EXTERNAL_SURVEYOR : 'm',
  SurveyorTypeKey.SELF : 'k',
  SurveyorTypeKey.REPORTING_MANAGER : 'r',
  SurveyorTypeKey.DIRECT_REPORT : 'b',
  SurveyorTypeKey.PEER : 'g',
  SurveyorTypeKey.INTERNAL_SURVEYOR : 'c',
  SurveyorTypeKey.EXTERNAL_SURVEYOR : 'm'  
}


def FormatTableColumnWidth(column, width):
  column.width = width
  for cell in column.cells:
    cell.width = width
    print(cell)


def FormatTableCellAlignText(cell, alignment):
  cell.paragraphs[0].alignment = alignment


def FormatTableCellFontProperties(cell, font_name, font_size, bold_flag, italic_flag):
  if len(cell.paragraphs[0].runs):
    run = cell.paragraphs[0].runs[0]
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold_flag
    run.font.italic = italic_flag
  else:
    run = cell.paragraphs[0].add_run()
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold_flag
    run.font.italic = italic_flag


def FormatTableColumnFontProperties(column, start_idx, font_name, font_size, bold_flag, italic_flag):
  for cell in column.cells[start_idx:]:
    print(f"cell.text={cell.text}")
    print(f"Number of paragraphs={len(cell.paragraphs)}")
    print(f"Number of runs in paragraphs[0]={len(cell.paragraphs[0].runs)}")
    if len(cell.paragraphs[0].runs):
      run = cell.paragraphs[0].runs[0]
      run.font.name = font_name
      run.font.size = font_size
      run.font.bold = bold_flag
      run.font.italic = italic_flag
    else:
      run = cell.paragraphs[0].add_run()
      run.font.name = font_name
      run.font.size = font_size
      run.font.bold = bold_flag
      run.font.italic = italic_flag


def FormatTableRowFontProperties(row, start_idx, font_name, font_size, bold_flag, italic_flag):
  for cell in row.cells[start_idx:]:
    if len(cell.paragraphs[0].runs):
      run = cell.paragraphs[0].runs[0]
      run.font.name = font_name
      run.font.size = font_size
      run.font.bold = bold_flag
      run.font.italic = italic_flag
    else:
      run = cell.paragraphs[0].add_run()
      run.font.name = font_name
      run.font.size = font_size
      run.font.bold = bold_flag
      run.font.italic = italic_flag

    
def FormatTableCellFontColor(cell, font_rgb_str):
  if len(cell.paragraphs[0].runs):
    run = cell.paragraphs[0].runs[0]
    run.font.color.rgb = RGBColor.from_string(font_rgb_str)
  else:
    run = cell.paragraphs[0].add_run()
    run.font.color.rgb = RGBColor.from_string(font_rgb_str) 


def FormatTableCellBackgroundColor(cell, fill_rgb):
  print("FormatTableCellBackgroundColor: Before Change:")
  print(cell._tc.xml)
  xmle = '<w:shd {} w:fill="{}" />'.format(nsdecls('w'), fill_rgb)
  shading_elem = parse_xml(xmle)
  cell._tc.get_or_add_tcPr().append(shading_elem)
  print("FormatTableCellBackgroundColor: After Change:")
  print(cell._tc.xml)

   
def FormatTableBorders(table):
  tbl = table._tbl
  for cell in tbl.iter_tcs():
    tcPr = cell.tcPr
    tcBorders = OxmlElement("w:tcBorders")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn('w:sz'), '4')
    top.set(qn('w:space'), '0')
    top.set(qn('w:color'), '4F81BD')
  
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn('w:sz'), '4')
    left.set(qn('w:space'), '3')
    left.set(qn('w:color'), '4F81BD')
  
    right = OxmlElement("w:right")
    right.set(qn("w:val"), "single")
    right.set(qn('w:sz'), '4')
    right.set(qn('w:space'), '3')
    right.set(qn('w:color'), '4F81BD')
  
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), '4F81BD')
  
    tcBorders.append(top)
    tcBorders.append(left)
    tcBorders.append(bottom)
    tcBorders.append(right)
    tcPr.append(tcBorders)


def CreateCustomDocument():
  doc = docx.Document()
  styles = doc.styles
  paragraph_styles = [ s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH ]
  print("---------Paragraph Styles:Before---------------")
  for style in paragraph_styles:
    print(f"\tStyle Name={style.name}, font={style.font.name}, size={style.font.size}, color={style.font.color.rgb}, type={type(style.font.color.rgb)}\n")

  style = doc.styles['Normal']
  font = style.font
  font.name = 'Arial'
  font.size = Pt(6)

  title_style = doc.styles['Title']
  title_font  = title_style.font
  title_font.name = 'Arial'
  title_font.size = Pt(14)
  title_font.color.rgb = docx.shared.RGBColor(0x00,0x00, 0x00)
  
  head1_style = doc.styles['Heading 1']
  head1_style.font.name = 'Ariel'
  head1_style.font.size = Pt(10)
  head1_style.font.color.rgb = docx.shared.RGBColor(0x00,0x00, 0x00)

  head2_style = doc.styles['Heading 2']
  head2_style.font.name = 'Ariel'
  head2_style.font.size = Pt(9)
  head2_style.font.color.rgb = docx.shared.RGBColor(0x00,0x00, 0x00)

  head3_style = doc.styles['Heading 3']
  head3_style.font.name = 'Ariel'
  head3_style.font.size = Pt(8)
  head3_style.font.color.rgb = docx.shared.RGBColor(0x00,0x00, 0x00)

  normal_style = doc.styles['Normal']
  normal_style.font.name = 'Verdana'
  normal_style.font.size = Pt(6)
  #normal_style.font.color.rgb = docx.shared.RGBColor(0x00,0x00, 0x00)
  
  body_text_style = doc.styles['Body Text']
  body_text_style.font.name = 'Verdana'
  body_text_style.font.size = Pt(6)
  #body_text_style.font.color.rgb = docx.shared.RGBColor(0x00,0x00, 0x00)
  
  no_spacing_style = doc.styles['No Spacing']
  no_spacing_style.font.name = 'Verdana'
  no_spacing_style.font.size = Pt(6)
  #no_spacing_style.font.color.rgb = docx.shared.RGBColor(0x00,0x00, 0x00)
  
  list_num_style = doc.styles['List Number']
  list_num_style.font.name = 'Verdana'
  list_num_style.font.size = Pt(6)
  #list_num_style.font.color.rgb = docx.shared.RGBColor(0x00,0x00, 0x00)
  
  list_bullet_style = doc.styles['List Bullet']
  list_bullet_style.font.name = 'Verdana'
  list_bullet_style.font.size = Pt(6)
  #list_bullet_style.font.color.rgb = docx.shared.RGBColor(0x00,0x00, 0x00)
  return doc
  
def GetOrderedSurveyorList(stype_map, stype_list):
  print("-------------------GetOrderedSurveyorList:[start]-------------------")
  stype_list.clear()
  for stype_key in SurveyorTypeKey:
    print(f"SurveyorTypeKey={stype_key}")
    if stype_key in stype_map:
      print(f"Adding SurveyorTypeKey={stype_key} to list")
      stype_list.append(stype_key)
  print("-------------------GetOrderedSurveyorList:[end]-------------------")    

  
def GetOpenAssignmnetCountForEmployeeInSurvey(db_con_obj, survey_id, eid):
  print("-------------------GetOpenAssignmnetCountForEmployeeInSurvey:[start]------------------------\n");
  print(f"Looking for open assignments in survey={survey_id} to evaluate employee=[{eid}] typeofeid={type(eid)}\n")
  open_assignments = 1

  con = db_con_obj.GetConnection()
  if not con:
    print("GetOpenAssignmnetCountForEmployeeInSurvey: Failed to open DB Connection")
    sys.exit(2)
  if con.is_connected():
    cursor = con.cursor()
    sql = "SELECT COUNT(*) FROM SurveyAssignment WHERE SurveyID={} AND EID='{}' AND ISNULL(ResponseDate)".format(survey_id, eid)
    print("using new count")
    try:
      cursor.execute(sql)
      print(f"Number of rows returned by cursor is={cursor.rowcount}\n")
      row = cursor.fetchone()
      open_assignments = row[0]
      print(f"Total sum returned={open_assignments}\n")
      if open_assignments:
        print(f"Number of open Assignments for Survey={survey_id} of Employee={eid} is={open_assignments}\n")
      else:
        print("No open Assignments for this user\n")
    except (mysql.connector.Error) as e:
      print(f"GetOpenAssignmnetCountForEmployeeInSurvey: Failed due to exception:{e}\n")
    finally:
      print("closing cursor\n")
      cursor.close()
  print("-------------------GetOpenAssignmnetCountForEmployeeInSurvey:[end]------------------------\n")
  #Error conditions return 1
  print(f"Returning Open Assignment count={open_assignments}")
  return open_assignments


def GetQuestionListIdOfSurvey(db_con_obj, survey_id):
  print("-------------------GetQuestionListIdOfSurvey:[start]------------------------\n")
  question_list_id = 0
  con = db_con_obj.GetConnection()
  try:
    survey_state_sql = f"SELECT QuestionaireID FROM Survey WHERE SurveyID = {survey_id}"
    cursor = con.cursor()
    cursor.execute(survey_state_sql)
    result_row = cursor.fetchone()
    if (result_row):
      question_list_id = result_row[0];
      print(f"Survey SurveyID={survey_id}, QuestionaireID={question_list_id}")
  except (mysql.connector.Error, mysql.connector.Warning) as e:
    print(f"GetQuestionListIdOfSurvey: Failed due to exception:{e}\n")
  print("-------------------GetQuestionListIdOfSurvey:[start]------------------------\n")
  return question_list_id
  

def DepartmentId2DepartmentName(db_con_obj, department_id):
  print("-------------------DepartmentId2DepartmentName:[start]------------------------\n");
  department_name = None
  con = db_con_obj.GetConnection()
  if not con:
    print("DepartmentId2DepartmentName: Failed to open DB Connection")
    sys.exit(2)
  if not con.is_connected():
    print("Database is not connected\n")
    sys.exit(2)
  sql= f"SELECT DepartmentName FROM Department WHERE DepartmentID={department_id}"
  try:
    cursor = con.cursor()
    cursor.execute(sql)
    row = cursor.fetchone()
    if row:
      print(f"The DepartmentName of DepartmentId={department_id} is={row[0]}\n")
      department_name = row[0]
  except (mysql.connector.Error, mysql.connector.Warning) as e:
    print(f"DepartmentId2DepartmentName: Failed due to exception:{e}\n")
  finally:
    cursor.close()  
  print("-------------------DepartmentId2DepartmentName:[end]------------------------\n");
  return department_name


def JobTitleId2JobTitleName(db_con_obj, job_title_id):
  print("-------------------JobTitleId2JobTitleName:[start]------------------------\n");
  job_title_name = None
  con = db_con_obj.GetConnection()
  if not con:
    print("JobTitleId2JobTitleName: Failed to open DB Connection")
    sys.exit(2)
  if not con.is_connected():
    print("Database is not connected\n")
    sys.exit(2)
  sql= f"SELECT JobTitleName FROM JobTitle WHERE JobTitleID={job_title_id}"
  try:
    cursor = con.cursor()
    cursor.execute(sql)
    row = cursor.fetchone()
    if row:
      print(f"The JobTitleName of JobTitleId={job_title_id} is={row[0]}\n")
      job_title_name = row[0]
  except (mysql.connector.Error, mysql.connector.Warning) as e:
    print(f"JobTitleId2JobTitleName: Failed due to exception:{e}\n")
  finally:
    cursor.close()  
  print("-------------------JobTitleId2JobTitleName:[end]------------------------\n");
  return job_title_name


def GetQuestionaireDetailsFromQuestionListId(db_con_obj, question_list_id, questionaire_details_map):
  print("-------------------------GetQuestionaireDetailsFromQuestionListId:[start]-------------------------");
  con = db_con_obj.GetConnection()
  if not con:
    print("GetQuestionaireDetailsFromQuestionListId: Failed to open DB Connection")
    sys.exit(2)
  questionaire_details_map.clear();
  category_list_sql = f"SELECT  CategoryID, CategoryName FROM QuestionCategoryTable WHERE QuestionListID ={question_list_id} ORDER BY CategoryID ASC";
  questionaire_details_map['QuestionListID'] = question_list_id
  questionaire_details_map['CategoryMap'] = OrderedDict()
  try:
    cursor = con.cursor()
    cursor.execute(category_list_sql)
    category_detail_rows = cursor.fetchall()
    if (not len(category_detail_rows)):
      print("Warning: GetQuestionaireDetails: Empty Question List returned by DB")
      return
    
    for category_details_row in category_detail_rows:
      category_id = category_details_row[0];
      category_name = category_details_row[1];
      questionaire_details_map['CategoryMap'][category_id] = {
        'CategoryName' : category_name,
        'QuestionMap' : OrderedDict()
        }
      print(f"Looking up questions for QuestionCategoryId ID:{question_list_id}, CategoryId={category_id}, CategoryName:{category_name}")
      question_list_sql = f"SELECT QuestionNumber, Question FROM QuestionsTable WHERE QuestionListID ={question_list_id} AND CategoryID ={category_id} ORDER BY QuestionNumber ASC"
      cursor.execute(question_list_sql)
      question_list_rows = cursor.fetchall()
      for question_details_row in question_list_rows:
        question_number = question_details_row[0]
        question = question_details_row[1]
        print(f"\tAdding question number={question_number}, Question={question}")
        questionaire_details_map['CategoryMap'][category_id]['QuestionMap'][question_number] = question
  except (mysql.connector.Error, mysql.connector.Warning) as e:
    print(f"GetQuestionaireDetailsFromQuestionListId: Failed due to exception:{e}\n")
  finally:
    cursor.close()
  print("-------------------GetQuestionaireDetailsFromQuestionListId:[end]------------------------");


def GetCategoryListFromQuestionaireDetails(questionaire_details_map, category_list):
  print("-------------------------GetCategoryListFromQuestionaireDetails:[start]-------------------------")
  category_list.clear()
  for category_id in questionaire_details_map['CategoryMap'].keys():
    print(f"Adding CategoryId={category_id}, CategoryName={questionaire_details_map['CategoryMap'][category_id]['CategoryName']}")
    category_list.append((category_id, questionaire_details_map['CategoryMap'][category_id]['CategoryName']))
  category_list.sort(key=lambda tup: tup[0])
  print(f"Total Number of categories found={len(category_list)}")
  print("-------------------------GetCategoryListFromQuestionaireDetails:[end]-------------------------")


def GetEmplolyeeDetails(db_con_obj, eid, employee_details_map):
  print("-------------------------GetEmplolyeeDetails:[start]-------------------------");
  con = db_con_obj.GetConnection()
  if not con:
    print("GetEmplolyeeDetails: Failed to open DB Connection")
    sys.exit(2)
  if not con.is_connected():
    print("Database is not connected\n")
    sys.exit(2)
  employee_details_map.clear()
  sql= "SELECT Name, BranchID, DepartmentID, JobTitleID, PhoneNumber, Email FROM Employee WHERE EID='{}'".format(eid)
  try:
    cursor = con.cursor()
    row_count = cursor.execute(sql)
    print(f"Number of rows returned={row_count}\n")
    employee_details = cursor.fetch()
    employee_details_map['EID'] = eid
    employee_details_map['Name'] = employee_details[0]
    employee_details_map['BranchID'] = employee_details[1]
    employee_details_map['Department'] = DepartmentId2DepartmentName(db_con_obj, employee_details[2])
    employee_details_map['JobTitle'] = JobTitleId2JobTitleName(db_con_obj, employee_details[3])
    employee_details_map['PhoneNumber'] = employee_details[4]
    employee_details_map['Email'] = employee_details[5]
    print(f"Employee Details: EID={employee_details_map['EID']}, ")
  except (mysql.connector.Error, mysql.connector.Warning) as e:
      print(f"GetEmplolyeeDetails: Failed due to exception:{e}\n")
  finally:
    cursor.close()  
  print("-------------------------GetEmplolyeeDetails:[end]-------------------------")
  
  
def radar_factory(num_vars, frame='circle'):
  """Create a radar chart with `num_vars` axes.

  This function creates a RadarAxes projection and registers it.

  Parameters
  ----------
    num_vars : int
      Number of variables for radar chart.
    frame : {'circle' | 'polygon'}
      Shape of frame surrounding axes.

  """
  # calculate evenly-spaced axis angles
  theta = np.linspace(0, 2 * np.pi, num_vars, endpoint=False)

  class RadarTransform(PolarAxes.PolarTransform):

    def transform_path_non_affine(self, path):
      # Paths with non-unit interpolation steps correspond to gridlines,
      # in which case we force interpolation (to defeat PolarTransform's
      # autoconversion to circular arcs).
      if path._interpolation_steps > 1:
        path = path.interpolated(num_vars)
      return Path(self.transform(path.vertices), path.codes)

  class RadarAxes(PolarAxes):

    name = 'radar'
        
    PolarTransform = RadarTransform

    def __init__(self, *args, **kwargs):
      super().__init__(*args, **kwargs)
      # rotate plot such that the first axis is at the top
      self.set_theta_zero_location('N')

    def fill(self, *args, closed=True, **kwargs):
      """Override fill so that line is closed by default"""
      return super().fill(closed=closed, *args, **kwargs)

    def plot(self, *args, **kwargs):
      """Override plot so that line is closed by default"""
      lines = super().plot(*args, **kwargs)
      for line in lines:
        self._close_line(line)

    def _close_line(self, line):
      x, y = line.get_data()
      # FIXME: markers at x[0], y[0] get doubled-up
      if x[0] != x[-1]:
        x = np.append(x, x[0])
        y = np.append(y, y[0])
        line.set_data(x, y)

    def set_varlabels(self, labels):
      print(f"set_labels={labels}")
      self.set_thetagrids(np.degrees(theta), labels)

    def _gen_axes_patch(self):
      # The Axes patch must be centered at (0.5, 0.5) and of radius 0.5
      # in axes coordinates.
      if frame == 'circle':
        print("_gen_axes_patch:Using frame=circle")
        return Circle((0.5, 0.5), 0.5)
      elif frame == 'polygon':
        print("_gen_axes_patch:Using frame=polygon")
        return RegularPolygon((0.5, 0.5), num_vars,
                              radius=.5, edgecolor="k")
      else:
        raise ValueError("unknown value for 'frame': %s" % frame)

    def draw(self, renderer):
      """ Draw. If frame is polygon, make gridlines polygon-shaped """
      if frame == 'polygon':
        print("draw:Using polygon")
        gridlines = self.yaxis.get_gridlines()
        for gl in gridlines:
          gl.get_path()._interpolation_steps = num_vars
      super().draw(renderer)

    def _gen_axes_spines(self):
      if frame == 'circle':
        return super()._gen_axes_spines()
      elif frame == 'polygon':
        # spine_type must be 'left'/'right'/'top'/'bottom'/'circle'.
        print(f"_gen_axes_spines:Using polygon, num_vars={num_vars}")
        spine = Spine(axes=self,
                      spine_type='circle',
                      path=Path.unit_regular_polygon(num_vars))
        # unit_regular_polygon gives a polygon of radius 1 centered at
        # (0, 0) but we want a polygon of radius 0.5 centered at (0.5,
        # 0.5) in axes coordinates.
        spine.set_transform(Affine2D().scale(.5).translate(.5, .5)
                            +self.transAxes)

        return {'polar': spine}
      else:
        raise ValueError("unknown value for 'frame': %s" % frame)

  register_projection(RadarAxes)
  return theta


def GetSurveyorInformationForAssignment(db_con_obj, survey_id, eid, eid_map, surveyor_type_map):
  print("-------------------GetSurveyorInformationForAssignment:[start]------------------------\n");
  con = db_con_obj.GetConnection()
  if not con:
    print("GetSurveyFeedbackForEmployeeBySurveyorType: Failed to open DB Connection")
    sys.exit(2)
  if not con.is_connected():
    print("Database is not connected\n")
    sys.exit(2)

  eid_map.clear()
  surveyor_type_map.clear()
  try:
    cursor = con.cursor()
    assignments_sql = "SELECT SurveyAssignment.SurveyorEID, SurveyAssignment.SurveyorType, Employee.Name,"
    assignments_sql += " Employee.DepartmentID, Employee.JobTitleID FROM SurveyAssignment INNER JOIN Employee on SurveyAssignment.SurveyorEID = Employee.EID"
    assignments_sql += " WHERE SurveyAssignment.SurveyID={} AND SurveyAssignment.EID='{}'".format(survey_id, eid)
    print(f"Executing SQL=[{assignments_sql}]\n")
    cursor.execute(assignments_sql)
    print(f"Number of rows returned:{cursor.rowcount}")
    survey_assignments = cursor.fetchall()
    for survey_assignment in survey_assignments:
      surveyor_eid = str(survey_assignment[0])
      surveyor_type_id = survey_assignment[1]
      surveyor_name = survey_assignment[2]
      surveyor_dept_id = survey_assignment[3]
      surveyor_job_title_id = survey_assignment[4]
      surveyor_type = SurveyorType2SurveyorTypeKey[SurveyorType(surveyor_type_id)]
      surveyor_dept = DepartmentId2DepartmentName(db_con_obj, surveyor_dept_id)
      surveyor_job_title = JobTitleId2JobTitleName(db_con_obj, surveyor_job_title_id)
      print(f"Adding Surveyor: EId={surveyor_eid}, TypeID={surveyor_type_id}, Type={surveyor_type}, Name={surveyor_name}, Department={surveyor_dept}, JobTitle={surveyor_job_title}\n")
      eid_map[surveyor_eid] = {'TypeId' : surveyor_type_id, 'Type' : surveyor_type,
                              'Name' : surveyor_name, 'Department' : surveyor_dept,
                              'JobTitle' : surveyor_job_title}
      if surveyor_type not in surveyor_type_map:
        surveyor_type_map[surveyor_type] = {}
      surveyor_type_map[surveyor_type][surveyor_eid] = {'Name' : surveyor_name, 'TypeId' : surveyor_type_id,
                                                        'Department' : surveyor_dept,
                                                        'JobTitle' : surveyor_job_title}
  except (mysql.connector.Error, mysql.connector.Warning) as e:
    print(f"GetSurveyFeedbackForEmployee: Failed due to mysql exception:{e}\n")
  finally:
    cursor.close()  
   
  print("-------------------GetSurveyorInformationForAssignment:[end]------------------------\n");

  
def GetSurveyFeedbackForEmployee(db_con_obj, survey_id, eid, ratings_list):
  print("-------------------GetSurveyFeedbackForEmployee:[start]------------------------\n");
  con = db_con_obj.GetConnection()
  if not con:
    print("GetSurveyFeedbackForEmployee: Failed to open DB Connection")
    sys.exit(1)
  if not con.is_connected():
    print("Database is not connected\n")
    sys.exit(2)
  try:
    cursor = con.cursor()
    sql = f"select SurveyorEID, SurveyorType, QuestionCategoryID, QuestionID, Rating "
    sql += "from SurveyFeedback where SurveyID={} and EID='{}'".format(survey_id, eid)
    print(f"Executing internal query:[{sql}]\n")
    cursor.execute(sql)
    print(f"Downloading Survey Feedback for Employee={eid} Survey={survey_id}\n")
    print(f"size of the ratings_list before clearing={len(ratings_list)}")
    ratings_list.clear()
    feedback_list = cursor.fetchall()
    for feedback_instance in feedback_list:
      surveyor_eid = feedback_instance[0]
      surveyor_type = feedback_instance[1]
      category_id = feedback_instance[2]
      question_id = feedback_instance[3]
      rating = feedback_instance[4]
      print(f"\tAdding: SurveyorEID={surveyor_eid}, SurveyorType={surveyor_type}, CategoryId={category_id}, QuestionId={question_id}, Rating={rating}\n")
      ratings_list.append((surveyor_eid, surveyor_type, category_id, question_id, rating))
    print(f"Number of ratings found for employee={eid} for survey={survey_id} is={len(ratings_list)}")    
  except (mysql.connector.Error, mysql.connector.Warning) as e:
      print(f"GetSurveyFeedbackForEmployee: Failed due to mysql exception:{e}\n")
  finally:
    cursor.close()  
  print("-------------------GetSurveyFeedbackForEmployee:[end]------------------------\n");


def ProcessData(ratings_array, data_map, questonaire_map):
  print("-------------------ProcessData:[start]------------------------\n");
  data_map.clear()
  print(f"----------Calculating Total Ratings------------------\n")
  data_map['TotalRatings'] = OrderedDict()
  data_map['DetailedRatings'] = OrderedDict()
  data_map['TotalRatings']['SelfAverage'] = np.round(np.average(ratings_array[ratings_array['type'] == SurveyorType.SELF]['rating']), 2)
  data_map['TotalRatings']['SelfMean'] = np.round(np.mean(ratings_array[ratings_array['type'] == SurveyorType.SELF]['rating']), 2)
  data_map['TotalRatings']['SelfStdDeviation'] = np.round(np.std(ratings_array[ratings_array['type'] == SurveyorType.SELF]['rating']), 2)  
  data_map['TotalRatings']['SelfSum'] = np.sum(ratings_array[ratings_array['type'] == SurveyorType.SELF]['rating'])
  data_map['TotalRatings']['SelfHighestRating'] = np.max(ratings_array[ratings_array['type'] == SurveyorType.SELF]['rating'])
  data_map['TotalRatings']['SelfLowestRating'] = np.min(ratings_array[ratings_array['type'] == SurveyorType.SELF]['rating'])
  data_map['TotalRatings']['SelfHighestRatingCount'] = np.count_nonzero((ratings_array['type'] == SurveyorType.SELF) & 
                                                                        (ratings_array['rating'] == data_map['TotalRatings']['SelfHighestRating'])) 
  data_map['TotalRatings']['SelfRatingsCount'] = ratings_array[ratings_array['type'] == SurveyorType.SELF].size
  print(f"""TotalRatings[Self Survey]: Average={data_map['TotalRatings']['SelfAverage']}, Mean={data_map['TotalRatings']['SelfMean']}, 
  Sum={data_map['TotalRatings']['SelfSum']}, RatingsCount={data_map['TotalRatings']['SelfRatingsCount']}""")
  data_map['TotalRatings']['OtherAverage'] = np.round(np.average(ratings_array[ratings_array['type'] != SurveyorType.SELF]['rating']), 2)
  data_map['TotalRatings']['OtherMean'] = np.round(np.mean(ratings_array[ratings_array['type'] != SurveyorType.SELF]['rating']), 2)
  data_map['TotalRatings']['OtherStdDeviation'] = np.round(np.std(ratings_array[ratings_array['type'] != SurveyorType.SELF]['rating']), 2)  
  data_map['TotalRatings']['OtherSum'] = np.sum(ratings_array[ratings_array['type'] != SurveyorType.SELF]['rating'])
  data_map['TotalRatings']['OtherHighestRating'] = np.max(ratings_array[ratings_array['type'] != SurveyorType.SELF]['rating'])
  data_map['TotalRatings']['OtherLowestRating'] = np.min(ratings_array[ratings_array['type'] != SurveyorType.SELF]['rating'])
  data_map['TotalRatings']['OtherHighestRatingCount'] = np.count_nonzero((ratings_array['type'] != SurveyorType.SELF) & 
                                                                        (ratings_array['rating'] == data_map['TotalRatings']['OtherHighestRating']))  
  data_map['TotalRatings']['OtherRatingsCount'] = ratings_array[ratings_array['type'] != SurveyorType.SELF].size
  print(f"""TotalRatings[Other Surveyors]: Average={data_map['TotalRatings']['OtherAverage']}, Mean={data_map['TotalRatings']['OtherMean']}, 
  Sum={data_map['TotalRatings']['OtherSum']}, RatingsCount={data_map['TotalRatings']['OtherRatingsCount']}""")
  

  data_map['DetailedRatings'][SurveyorTypeKey.SELF]={}
  if ratings_array[ratings_array['type'] == SurveyorType.REPORTING_MANAGER].size:
    print(f"Creating DetailedRatings map for surveyor type={SurveyorTypeKey.REPORTING_MANAGER}\n")
    data_map['DetailedRatings'][SurveyorTypeKey.REPORTING_MANAGER]={}
  if ratings_array[ratings_array['type'] == SurveyorType.DIRECT_REPORT].size:
    print(f"Creating DetailedRatings map for surveyor type={SurveyorTypeKey.DIRECT_REPORT}\n")
    data_map['DetailedRatings'][SurveyorTypeKey.DIRECT_REPORT]={}
  if ratings_array[ratings_array['type'] == SurveyorType.PEER].size:
    data_map['DetailedRatings'][SurveyorTypeKey.PEER]={}
  if ratings_array[ratings_array['type'] == SurveyorType.INTERNAL_SURVEYOR].size:
    data_map['DetailedRatings'][SurveyorTypeKey.INTERNAL_SURVEYOR]={}
  if ratings_array[ratings_array['type'] == SurveyorType.EXTERNAL_SURVEYOR].size:
    data_map['DetailedRatings'][SurveyorTypeKey.EXTERNAL_SURVEYOR]={}
  for surveyor_type_key in data_map['DetailedRatings']:
    surveyor_type_id = SurveyorTypeKey2SurveyorType[surveyor_type_key]
    print(f"--------------------Extracting data for surveyor_type={surveyor_type_key}, id={surveyor_type_id}--------------------\n")
    data_map['DetailedRatings'][surveyor_type_key]['RatingsAveragePerCategory'] = []
    data_map['DetailedRatings'][surveyor_type_key]['RatingsMeanPerCategory'] = []
    data_map['DetailedRatings'][surveyor_type_key]['RatingsStdDeviationPerCategory'] = []
    data_map['DetailedRatings'][surveyor_type_key]['RatingsMinimum'] = []
    data_map['DetailedRatings'][surveyor_type_key]['RatingsMaximum'] = []
    for category_id in questonaire_map['CategoryMap']:
      print(f"Extracting Ratings averages for category_id={category_id}")
      avg_rating = np.round(np.average(ratings_array[(ratings_array['type'] == surveyor_type_id) & (ratings_array['category'] == category_id)]['rating']), 2)
      mean_rating = np.round(np.mean(ratings_array[(ratings_array['type'] == surveyor_type_id) & (ratings_array['category'] == category_id)]['rating']), 2)
      std_deviation = np.round(np.std(ratings_array[(ratings_array['type'] == surveyor_type_id) & (ratings_array['category'] == category_id)]['rating']), 2)
      min_rating = np.min(ratings_array[(ratings_array['type'] == surveyor_type_id) & (ratings_array['category'] == category_id)]['rating'])
      max_rating = np.max(ratings_array[(ratings_array['type'] == surveyor_type_id) & (ratings_array['category'] == category_id)]['rating'])
      print(f"Surveyor Type={surveyor_type_key}, Category Id={category_id}, Category Name={questonaire_map['CategoryMap'][category_id]['CategoryName']}, Average Rating={avg_rating}\n")
      print(f"Surveyor Type={surveyor_type_key}, Category Id={category_id}, Category Name={questonaire_map['CategoryMap'][category_id]['CategoryName']}, Mean Rating={mean_rating}\n")
      print(f"Surveyor Type={surveyor_type_key}, Category Id={category_id}, Category Name={questonaire_map['CategoryMap'][category_id]['CategoryName']}, Std Deviation={std_deviation}\n")
      print(f"Surveyor Type={surveyor_type_key}, Category Id={category_id}, Category Name={questonaire_map['CategoryMap'][category_id]['CategoryName']}, Minimum Rating={min_rating}\n")
      print(f"Surveyor Type={surveyor_type_key}, Category Id={category_id}, Category Name={questonaire_map['CategoryMap'][category_id]['CategoryName']}, Maximum Rating={max_rating}\n")
      data_map['DetailedRatings'][surveyor_type_key]['RatingsAveragePerCategory'].append(avg_rating)
      data_map['DetailedRatings'][surveyor_type_key]['RatingsMeanPerCategory'].append(mean_rating)
      data_map['DetailedRatings'][surveyor_type_key]['RatingsStdDeviationPerCategory'].append(std_deviation)
      data_map['DetailedRatings'][surveyor_type_key]['RatingsMinimum'].append(min_rating)
      data_map['DetailedRatings'][surveyor_type_key]['RatingsMaximum'].append(max_rating)
  print(data_map)
  print("-------------------ProcessData:[end]------------------------\n");


def PlotScatterGraphOfCategoryLevelFeedbackData(survey_id, eid, seid_2_info_map, category_name, question_list,
                                                  question_id_list, ratings_avg_map, doc):
  fig = plt.figure(figsize=(8.5, 5), layout='constrained')
  q_text = ""
  for qid in question_id_list:
    q_text += f"{qid}. {question_list[qid - 1]}\n"
  print(f"Using text[{q_text}] in the graph")
  #fig = plt.figure(figsize=(8.5, 5))
  ax = fig.subplots()
  #plt.text(5, 5, q_text, fontsize = 10)
  for surveyor_type_key in ratings_avg_map:
    plt.plot(question_id_list, ratings_avg_map[surveyor_type_key], label = SurveyorTypeKey2Label[surveyor_type_key],
             color = ColorMap[surveyor_type_key])
  plt.xlabel("Question Id List")
  plt.ylabel("Average Feedback Rating") 
  employee_name = seid_2_info_map[eid]['Name']
  plt.title(f"Survey ID:{survey_id} Average Feedback Received by employee: {employee_name} for Questions in Category: {category_name}")
  plt.legend()
  image_file = "survey_" + survey_id + employee_name.replace(" ", "_") + "Category_" + category_name.replace(" ", "_") + "_Questions_fb_scatter_graph.jpg"
  plt.savefig(image_file, bbox_inches='tight')
  doc.add_picture(image_file, width=Pt(300))
  if DisplayGraphsFlag:
    plt.show()
  if not KeepTempsFlag:
    os.remove(image_file)


def PlotRadarGraphsOfCategoryLevelFeedbackData(survey_id, eid, seid_2_info_map, category_name, question_list, 
                                               question_id_list, ratings_avg_map, doc):
  question_label_list = [str(qid) for qid in question_id_list]
  print("-------------PlotRadarGraphsOfCategoryLevelFeedbackData-----------------------")
  print(f"Question_id_list={question_id_list}, Question_list={question_list}")
  print(f"Question_label_list={question_label_list}")
  N = len(question_id_list)
  theta = radar_factory(N, frame='polygon')
  fig, (ax_cat_plain, ax_cat_filled) = plt.subplots(1, 2, figsize=(9, 9), layout='constrained', subplot_kw=dict(projection='radar'))
  fig.subplots_adjust(wspace=0.25, hspace=0.20, top=0.85, bottom=0.05)
  ax_cat_plain.set_title(f"Radar Graph", weight='bold', size='medium', position=(0.5, 1.1),
                         horizontalalignment='center', verticalalignment='center') 
  ax_cat_filled.set_title(f"Filled Radar Graph", weight='bold', size='medium', position=(0.5, 1.1),
                          horizontalalignment='center', verticalalignment='center')
  fig.suptitle(f"SurveyId:{survey_id} Rating averages for Category: {category_name}") 
  legend_strs = []
  for surveyor_type in ratings_avg_map:
    print(f"Plotting radar graph for surveyor type={surveyor_type}")
    color = ColorMap[surveyor_type]
    print(f"Ratings Average list={ratings_avg_map[surveyor_type]}")
    ax_cat_plain.plot(theta, ratings_avg_map[surveyor_type], color=color)
    ax_cat_filled.plot(theta, ratings_avg_map[surveyor_type], color=color)
    ax_cat_filled.fill(theta, ratings_avg_map[surveyor_type], facecolor=color, alpha=0.25, label='_nolegend_')
    print(f"Addomg ")
    legend_strs.append(SurveyorTypeKey2Label[surveyor_type])
  ax_cat_plain.set_varlabels(question_label_list) 
  ax_cat_filled.set_varlabels(question_label_list)
  print(f"Using legend_strs={legend_strs}") 
  ax_cat_plain.legend(legend_strs, loc=(0.9, .95), labelspacing=0.1, fontsize='small')
  employee_name = seid_2_info_map[eid]['Name']
  image_file = image_file = "survey_" + survey_id + employee_name.replace(" ", "_") + "Category_" + category_name.replace(" ", "_") + "_Questions_fb_radar_graph.jpg"
  plt.savefig(image_file, bbox_inches='tight')
  doc.add_picture(image_file, width=Pt(300))
  if DisplayGraphsFlag:
    plt.show()
  if not KeepTempsFlag:
    os.remove(image_file)

#This function directly analyzes data in the numpy array feedback_list
def WriteCategoryWiseReport(survey_id, eid, seid_2_info_map, ratings_array, surveyor_types_list, category_list,
                            questionaire_map, doc, sec_lst):
  heading_level = len(sec_lst)
  desc_str = f"""This section details the analytical data obtained by analyzing the feedback received for each question
  listed under each Category related to the Question List for the Survey Id {survey_id} conducted on the employee 
  {seid_2_info_map[eid]['Name']}"""  
  run = doc.add_paragraph().add_run(desc_str)
  run.font.name = "Verdana"
  run.font.size = Pt(6)

  for stype in surveyor_types_list:
    print(f"WriteCategoryWiseReport: surveyor_type={stype}")
  for category_id in questionaire_map['CategoryMap']:
    question_list = []
    question_id_list = []
    ratings_avg_map = OrderedDict([(stype, []) for stype in surveyor_types_list])
    ratings_mean_map = OrderedDict([(stype, []) for stype in surveyor_types_list])
    ratings_stddev_map = OrderedDict([(stype, []) for stype in surveyor_types_list])
    ratings_max_map = OrderedDict([(stype, []) for stype in surveyor_types_list])
    ratings_min_map = OrderedDict([(stype, []) for stype in surveyor_types_list])
    
    category_name = questionaire_map['CategoryMap'][category_id]['CategoryName']
    print(f"CategoryId={category_id}, Name={category_name}")
    for qid in questionaire_map['CategoryMap'][category_id]['QuestionMap']:
      print(f"QId={qid}, Question={questionaire_map['CategoryMap'][category_id]['QuestionMap'][qid]}")
      question_list.append(questionaire_map['CategoryMap'][category_id]['QuestionMap'][qid])
      question_id_list.append(qid)
      for stype in surveyor_types_list:
        stype_id = SurveyorTypeKey2SurveyorType[stype]
        print(f"Processing data for SurveyorType={stype_id} question id={qid}\n")
        avg_rating4q = np.round(np.average(ratings_array[(ratings_array['type'] == stype_id) & (ratings_array['category'] == category_id) & (ratings_array['question_number'] == qid)]['rating']), 2)
        mean_rating4q = np.round(np.mean(ratings_array[(ratings_array['type'] == stype_id) & (ratings_array['category'] == category_id) & (ratings_array['question_number'] == qid)]['rating']), 2)
        std_deviation4q = np.round(np.std(ratings_array[(ratings_array['type'] == stype_id) & (ratings_array['category'] == category_id) & (ratings_array['question_number'] == qid)]['rating']), 2)
        max_rating4q = np.max(ratings_array[(ratings_array['type'] == stype_id) & (ratings_array['category'] == category_id) & (ratings_array['question_number'] == qid)]['rating'])
        min_rating4q = np.min(ratings_array[(ratings_array['type'] == stype_id) & (ratings_array['category'] == category_id) & (ratings_array['question_number'] == qid)]['rating'])
        print(f"SurveyorTypeId={stype_id}, Category Id={category_id}, Question Id={qid}, Average Rating={avg_rating4q}, Mean Rating={mean_rating4q}, Min Rating={min_rating4q}, Max Rating={max_rating4q}\n")
        ratings_avg_map[stype].append(avg_rating4q)
        ratings_mean_map[stype].append(mean_rating4q)
        ratings_stddev_map[stype].append(std_deviation4q)
        ratings_max_map[stype].append(max_rating4q)
        ratings_min_map[stype].append(min_rating4q)
    print(f"SurveyorType={stype}, QAvg={ratings_avg_map[stype]}, MeanAvg={ratings_mean_map[stype]}, StdDev={ratings_stddev_map[stype]}, Min={ratings_max_map[stype]}, Max={ratings_min_map[stype]}")
    sec_lst[-1] = str(int(sec_lst[-1]) + 1)
    doc.add_heading(f"{'.'.join(sec_lst)} Detailed Analysis of Category:{category_name}", level= heading_level)
    run = doc.add_paragraph().add_run(f"Following is the summary of analytical data obtained for each question listed under the category {category_name}")
    run.font.name = "Verdana"
    run.font.size = Pt(6)
    doc.add_heading(f"{'.'.join(sec_lst) + '.1'} Analytic Data Table:{category_name}", level= heading_level + 1)
        
    row = 0
    header_rows = 3
    data_table = doc.add_table(rows=len(surveyor_types_list) * 5 + header_rows, cols = (len(question_list) + 1))
    print(f"len(surveyor_types_list)={len(surveyor_types_list)}, len(surveyor_types_list)*5={len(surveyor_types_list) * 5}, len(surveyor_types_list) * 5 + 2={len(surveyor_types_list) * 5 + 2}")
    print(f"data table size: rows={len(data_table.rows)}, columns={len(data_table.columns)}")
    data_table.style = 'Medium Grid 2 Accent 1'
    data_table.autofit = False
    data_table.allow_autofit = False
    data_table.cell(row, 1).merge(data_table.cell(row, len(question_list)))
    data_table.cell(row, 1).text = "Questions"
    FormatTableBorders(data_table)
    row += 1    
    hdr1_cells = data_table.rows[row].cells
    hdr2_cells = data_table.rows[row +1].cells
    hdr2_cells[0].text = "Question Id"
    for col in range(0, len(question_list)):
      print(f"Row={row},Col={col + 1}, question={question_list[col]}")
      print(f"Row={row + 1},Col={col + 1}, question id={question_id_list[col]}")
      hdr1_cells[col + 1].text = question_list[col]
      hdr2_cells[col + 1].text = str(question_id_list[col])
    row += 2
    start_row = row
    print(f"Start Row={start_row}")
    for idx in range(0, len(surveyor_types_list)):
      row = start_row + idx
      stype = surveyor_types_list[idx]
      print(f"Populating table for surveyor type {stype}, row={row}")
      col = 0
      print(f"Setting Avg surveyor type={stype}, row={row}, col={col}")
      data_table.cell(row, col).text = stype + '-Avg'
      for val in ratings_avg_map[stype]:
        col += 1
        data_table.cell(row, col).text = str(val)
      row += len(surveyor_types_list)
      col = 0
      print(f"Setting Mean surveyor type={stype}, row={row}, col={col}")
      data_table.cell(row, col).text = stype + '-Mean'     
      for val in ratings_mean_map[stype]:
        col += 1
        data_table.cell(row, col).text = str(val)
      row += len(surveyor_types_list)
      col = 0
      print(f"Setting StdDev surveyor type={stype}, row={row}, col={col}")
      data_table.cell(row, col).text = stype + '-StdDev'     
      for val in ratings_stddev_map[stype]:
        col += 1
        data_table.cell(row, col).text = str(val)
      row += len(surveyor_types_list)
      col = 0
      print(f"Setting Max surveyor type={stype}, row={row}, col={col}")
      data_table.cell(row, col).text = stype + '-MaxRating'     
      for val in ratings_max_map[stype]:
        col += 1
        data_table.cell(row, col).text = str(val)
      row += len(surveyor_types_list)
      col = 0
      print(f"Setting Min surveyor type={stype}, row={row}, col={col}")
      data_table.cell(row, col).text = stype + '-MinRating'     
      for val in ratings_min_map[stype]:
        col += 1
        data_table.cell(row, col).text = str(val)
    FormatTableCellAlignText(data_table.cell(0,1), WD_ALIGN_PARAGRAPH.CENTER)
    FormatTableCellFontProperties(data_table.cell(0,1), "Verdana", Pt(6), True, False)
    FormatTableRowFontProperties(data_table.rows[1], 1, "Verdana", Pt(6), False, True)
    FormatTableColumnFontProperties(data_table.columns[0], 1, "Verdana", Pt(6), False, False)
    doc.add_heading(f"{'.'.join(sec_lst) + '.2'} List of Questions for Category: {category_name}", level= heading_level + 1)
    #Using a customized paragraph to implement a numbered list as docx does not provde a way to reset the item number
    for i in range(0, len(question_list)):
      para = doc.add_paragraph()
      run = para.add_run(f"\t{i+1}. {question_list[i]}")
      para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
      run.font.size = Pt(6)
      run.italic =True    
    doc.add_heading(f"{'.'.join(sec_lst) + '.3'} Scatter Graph of Analyzed Data for Category: {category_name}", level= heading_level + 1)
    PlotScatterGraphOfCategoryLevelFeedbackData(survey_id, eid, seid_2_info_map, category_name, question_list, question_id_list, ratings_avg_map, doc)    
    doc.add_heading(f"{'.'.join(sec_lst) + '.4'} Radar Graphs of Analyzed Data for Category: {category_name}", level= heading_level + 1)
    PlotRadarGraphsOfCategoryLevelFeedbackData(survey_id, eid, seid_2_info_map, category_name, question_list, question_id_list, ratings_avg_map, doc)    

plt.style.use('_mpl-gallery')

    
def PlotBarGraphOfRatingsSummary(survey_id, eid, seid_2_info_map, category_names, data_map, doc):
  fig = plt.figure(figsize=(8.5, 1.5), layout='constrained')
  ax = fig.subplots()
  surveyor_type_labels = ('Self', 'Other\nSurveyors')
  #bar_colors = ['tab:red', 'tab:blue']
  bar_colors = ['#66CC00', '#FF6666']
  employee_name = seid_2_info_map[eid]['Name'] 
  y_pos = np.arange(len(surveyor_type_labels))
  print(f"y_pos={y_pos}\n")
  feed_back_avgs = [data_map['TotalRatings']['SelfAverage'], data_map['TotalRatings']['OtherAverage']]
  ax.barh(y_pos, feed_back_avgs, 0.75, color = bar_colors)
  ax.set_yticks(y_pos, labels=surveyor_type_labels)
  ax.invert_yaxis()
  ax.set_xlabel('Average Feedback Ratings')  
  ax.set_title(f"CombinedAverage Feedback Ratings for {employee_name} from All Surveyor Types")  
  image_file = "survey_" + survey_id + employee_name.replace(" ", "_") + "ratings_summary_bar_graph.jpg"
  print(f"Saving image to file:{image_file}")
  plt.savefig(image_file, bbox_inches='tight')
  doc.add_picture(image_file, width=Pt(400))
  if DisplayGraphsFlag:
    plt.show()
  if not KeepTempsFlag:
    os.remove(image_file)


def PlotPieChartOfRatingsSummaryBreakdown(survey_id, eid, seid_2_info_map, ratings_array, doc):
  employee_name = seid_2_info_map[eid]['Name']
  self_ratings_array = ratings_array[(ratings_array['type'] == SurveyorType.SELF)]
  surveyor_ratings_array = ratings_array[(ratings_array['type'] != SurveyorType.SELF)]
  print(f"Number of Self Ratings={self_ratings_array.size}, Number of Surveyor Ratings={surveyor_ratings_array.size}")
  self_rating_levels, self_rating_counts = np.unique(self_ratings_array['rating'], return_counts=True)
  print(f"Self rating levels={self_rating_levels}, number of instances={self_rating_counts}")
  self_rating_pcts = np.around(self_rating_counts * 100 / self_ratings_array.size, 3)
  print(f"self_counts_pct={self_rating_pcts}")
  # self_rating_pct_map = OrderedDict(zip(self_rating_levels, self_rating_pcts))
  # print(f"Self Rating percentages map:{self_rating_pct_map}")
  surveyor_rating_levels, surveyor_rating_counts = np.unique(surveyor_ratings_array['rating'], return_counts=True)
  print(f"Surveyor rating levels={surveyor_rating_levels}, number of entries={surveyor_rating_counts}")
  surveyor_rating_pcts = np.around(surveyor_rating_counts * 100 / surveyor_ratings_array.size,3)
  print(f"survyeyor rating percentages={surveyor_rating_pcts}") 
  # surveyor_rating_pct_map = OrderedDict(zip(surveyor_rating_levels, surveyor_rating_pcts))
  # print(f"Surveyor Rating percentages map:{surveyor_rating_pct_map}")
  self_rating_labels_num = [str(r) for r in self_rating_levels]
  self_rating_labels_str = [RatingLevel2Label[r] for r in self_rating_levels]  
  surveyor_rating_labels_num = [str(r) for r in surveyor_rating_levels]
  surveyor_rating_labels_str = [RatingLevel2Label[r] for r in surveyor_rating_levels]  
  fig, (ax_self_ratings, ax_surveyor_ratings) = plt.subplots(1, 2, figsize=(10, 10), layout='constrained')
  fig.subplots_adjust(wspace=0.25, hspace=0.20, top=0.85, bottom=0.05)
  fig.suptitle(f"360 Survey Id {survey_id} of Employee {employee_name} Ratings Breakdown comparison with Surveyors")
  ax_self_ratings.set_title(f"Self Ratings Breakdown", weight='bold', size='medium', position=(0.5, 1.1),
                         horizontalalignment='center', verticalalignment='center') 
  ax_surveyor_ratings.set_title(f"Surveyor Ratings Breakdown", weight='bold', size='medium', position=(0.5, 1.1),
                         horizontalalignment='center', verticalalignment='center') 
  ax_self_ratings.pie(self_rating_counts, labels=self_rating_labels_num, startangle = 90, autopct="%1.1f%%")
  ax_surveyor_ratings.pie(surveyor_rating_counts, labels=surveyor_rating_labels_num, startangle = 90, autopct="%1.1f%%")
  image_file = "survey_" + survey_id + employee_name.replace(" ", "_") + "ratings_breakdown_pie_chart.jpg"
  plt.savefig(image_file, bbox_inches='tight')
  doc.add_picture(image_file, width=Pt(400))
  if DisplayGraphsFlag:
    plt.show()
  if not KeepTempsFlag:
    os.remove(image_file)

                           
def WriteRatingsSummaryReport(survey_id, eid, seid_2_info_map, category_names, ratings_array, data_map, doc, sec_lst):
  sec_lst[-1] = str(int(sec_lst[-1]) + 1)
  heading_level = len(sec_lst)
  sect_str = '.'.join(sec_lst)
  doc.add_heading("{} Combined Feedback Rating Averages".format(sect_str, level=heading_level))
  doc.add_paragraph("Following is the summary of combined feedback received across all the Question List Categories")
  headers = ("Surveyor Type", "Average", "Mean", "Std Deviation", "Sum", "Highest Rating", "Highest Rating\nCount", "Lowest Rating", "Total Ratings")
  ratings_summary_table = doc.add_table(rows=3, cols=len(headers))
  ratings_summary_table.style = 'Medium Grid 2 Accent 1'
  ratings_summary_table.autofit = False
  ratings_summary_table.allow_autofit = False
  FormatTableBorders(ratings_summary_table)
  hdr_cells = ratings_summary_table.rows[0].cells
  for hdr_idx in range(0, len(headers)):
    hdr_cells[hdr_idx].text = headers[hdr_idx]
  ratings_summary_table.cell(1,0).text = "Self"
  ratings_summary_table.cell(1,1).text = str(data_map['TotalRatings']['SelfAverage'])
  ratings_summary_table.cell(1,2).text = str(data_map['TotalRatings']['SelfMean'])
  ratings_summary_table.cell(1,3).text = str(data_map['TotalRatings']['SelfStdDeviation'])
  ratings_summary_table.cell(1,4).text = str(data_map['TotalRatings']['SelfSum'])
  ratings_summary_table.cell(1,5).text = str(data_map['TotalRatings']['SelfHighestRating'])
  ratings_summary_table.cell(1,6).text = str(data_map['TotalRatings']['SelfHighestRatingCount'])
  ratings_summary_table.cell(1,7).text = str(data_map['TotalRatings']['SelfLowestRating'])
  ratings_summary_table.cell(1,8).text = str(data_map['TotalRatings']['SelfRatingsCount'])
  ratings_summary_table.cell(2,0).text = "Other Surveyors"
  ratings_summary_table.cell(2,1).text = str(data_map['TotalRatings']['OtherAverage'])
  ratings_summary_table.cell(2,2).text = str(data_map['TotalRatings']['OtherMean'])
  ratings_summary_table.cell(2,3).text = str(data_map['TotalRatings']['OtherStdDeviation'])
  ratings_summary_table.cell(2,4).text = str(data_map['TotalRatings']['OtherSum'])
  ratings_summary_table.cell(2,5).text = str(data_map['TotalRatings']['OtherHighestRating'])
  ratings_summary_table.cell(2,6).text = str(data_map['TotalRatings']['OtherHighestRatingCount'])
  ratings_summary_table.cell(2,7).text = str(data_map['TotalRatings']['OtherLowestRating'])
  ratings_summary_table.cell(2,8).text = str(data_map['TotalRatings']['OtherRatingsCount'])  
  FormatTableColumnFontProperties(ratings_summary_table.columns[0], 1, 'Verdana', Pt(6), False, False)
  sec_lst[-1] = str(int(sec_lst[-1]) + 1)
  sect_str = '.'.join(sec_lst)  
  doc.add_heading(f"{sect_str} Bar chart of the  Combined Feedback Rating Averages for Employee {seid_2_info_map[eid]['Name']} for Survey {survey_id}",
                                                                                                                level=heading_level)  
  PlotBarGraphOfRatingsSummary(survey_id, eid, seid_2_info_map, category_names, data_map, doc)
  sec_lst[-1] = str(int(sec_lst[-1]) + 1)
  sect_str = '.'.join(sec_lst)  
  doc.add_heading(f"{sect_str} Pie Chart of Ratings Breakdown between the Employee and Surveyors")
  PlotPieChartOfRatingsSummaryBreakdown(survey_id, eid, seid_2_info_map, ratings_array, doc)


def PlotScatterGraphpOfFeedbackAveragesBetweenSelfAndSurveyorType(survey_id, eid, surveyor_type_key, seid_2_info_map, category_names, data_map, doc):  
  print(f"Creating Scatter graph of Feedback received for employee:{eid} from Surveyor Type:{surveyor_type_key}")
  print(f"Using category list={category_names}\n")
  fig = plt.figure(figsize=(8.5, 5), layout='constrained')
  ax = fig.subplots()
  plt.plot(category_names, data_map['DetailedRatings']['Self']['RatingsAveragePerCategory'], label = SurveyorTypeKey2Label['Self'])
  plt.plot(category_names, data_map['DetailedRatings'][surveyor_type_key]['RatingsAveragePerCategory'],
           label = SurveyorTypeKey2Label[surveyor_type_key], color = ColorMap[surveyor_type_key])
  plt.xlabel("Evaluation Category")
  plt.ylabel("Average Feedback Rating") 
  employee_name = seid_2_info_map[eid]['Name'] 
  plt.title(f"Survey ID: {survey_id} Average Feedback Received from Surveyor Type: {SurveyorTypeKey2Label[surveyor_type_key]} for employee: {employee_name}")
  plt.legend()
  image_file = "survey_" + survey_id + employee_name.replace(" ", "_") + "_from_surveyor_type_" + surveyor_type_key  + "_scatter_graph.jpg"
  plt.savefig(image_file, bbox_inches='tight')
  doc.add_picture(image_file, width=Pt(300))
  if DisplayGraphsFlag:
    plt.show()
  if not KeepTempsFlag:
    os.remove(image_file)


def PlotRadarGraphOfFeedbackAveragesBetweenSelfAndSurveyorType(survey_id, eid, surveyor_type_key, seid_2_info_map, category_names, data_map, doc):
  print(f"Creating Radar graph of Feedback received for employee:{eid} from Surveyor Type:{surveyor_type_key}")
  N = len(category_names)
  theta = radar_factory(N, frame='polygon')
  print(f"Using N={N}, theta={theta}")
  fig, ax = plt.subplots(1, 1, figsize=(8, 8), layout='constrained', subplot_kw=dict(projection='radar'))
  fig.subplots_adjust(wspace=0.25, hspace=0.20, top=0.85, bottom=0.05)
  ax.set_title(f"Radar Graph of Feedback Received from Self and Surveyor Type:{SurveyorTypeKey2Label[surveyor_type_key]}", weight='bold', size='medium', position=(0.5, 1.1),
              horizontalalignment='center', verticalalignment='center')
  employee_name = seid_2_info_map[eid]['Name']    
  legend_strs = []
  ax.plot(theta, data_map['DetailedRatings']['Self']['RatingsAveragePerCategory'], color=ColorMap['Self'])
  legend_strs.append(SurveyorTypeKey2Label['Self'])
  ax.plot(theta, data_map['DetailedRatings'][surveyor_type_key]['RatingsAveragePerCategory'], color=ColorMap[surveyor_type_key])
  legend_strs.append(SurveyorTypeKey2Label[surveyor_type_key])  
  ax.set_varlabels(category_names)
  ax.legend(legend_strs, loc=(0.9, .95), labelspacing=0.1, fontsize='small')
  fig.suptitle(f"Survey ID: {survey_id} Average Feedback Received from Surveyor Type: {SurveyorTypeKey2Label[surveyor_type_key]} for employee: {employee_name}")
  image_file = "survey_" + survey_id + employee_name.replace(" ", "_") + "_from_surveyor_type_" + surveyor_type_key + "_radar_graph.jpg"
  print(f"Saving image to file:{image_file}")
  plt.savefig(image_file, bbox_inches='tight')
  doc.add_picture(image_file, width=Pt(300))
  if DisplayGraphsFlag:
    plt.show()
  if not KeepTempsFlag:
    os.remove(image_file)


def PlotFilledRadarGraphpOfFeedbackAveragesBetweenSelfAndSurveyorType(survey_id, eid, surveyor_type_key, seid_2_info_map, category_names, data_map, doc):
  print(f"Creating Filled Radar graph of Feedback received for employee:{eid} from Surveyor Type:{surveyor_type_key}")
  N = len(category_names)
  theta = radar_factory(N, frame='polygon')
  print(f"Using N={N}, theta={theta}")
  fig, ax = plt.subplots(1, 1, figsize=(8, 8), layout='constrained', subplot_kw=dict(projection='radar'))
  fig.subplots_adjust(wspace=0.25, hspace=0.20, top=0.85, bottom=0.05)
  ax.set_title(f"Radar Graph of Feedback Received from Self and Surveyor Type:{SurveyorTypeKey2Label[surveyor_type_key]}", weight='bold', size='medium', position=(0.5, 1.1),
              horizontalalignment='center', verticalalignment='center')    
  employee_name = seid_2_info_map[eid]['Name']
  legend_strs = []
  color = ColorMap['Self']
  ax.plot(theta, data_map['DetailedRatings']['Self']['RatingsAveragePerCategory'], color=color)
  ax.fill(theta, data_map['DetailedRatings']['Self']['RatingsAveragePerCategory'], facecolor=color, alpha=0.25, label='_nolegend_')
  legend_strs.append(SurveyorTypeKey2Label['Self'])
  color = ColorMap[surveyor_type_key]
  ax.plot(theta, data_map['DetailedRatings'][surveyor_type_key]['RatingsAveragePerCategory'], color=color)
  ax.fill(theta, data_map['DetailedRatings'][surveyor_type_key]['RatingsAveragePerCategory'], facecolor=color, alpha=0.25, label='_nolegend_')
  legend_strs.append(SurveyorTypeKey2Label[surveyor_type_key])
  ax.set_varlabels(category_names)
  ax.legend(legend_strs, loc=(0.9, .95), labelspacing=0.1, fontsize='small')
  fig.suptitle(f"Survey ID: {survey_id} Average Feedback Received from from Surveyor Type: {SurveyorTypeKey2Label[surveyor_type_key]} for employee: {employee_name}")
  image_file = "survey_" + survey_id + employee_name.replace(" ", "_") + "_fromSurveyorTYPE_" + surveyor_type_key + "_filled_radar_graph.jpg"
  print(f"Saving image to file:{image_file}")
  plt.savefig(image_file, bbox_inches='tight')
  doc.add_picture(image_file, width=Pt(300))
  if DisplayGraphsFlag:
    plt.show()
  if not KeepTempsFlag:
    os.remove(image_file)


def WriteFeedbackComparisonReport(survey_id, eid, seid_2_info_map, category_names, data_map, doc, sec_lst):
  heading_level = len(sec_lst)
  for surveyor_type_key in data_map['DetailedRatings'].keys():
    if 'Self' == surveyor_type_key:
      continue
    sec_lst[-1] = str(int(sec_lst[-1]) + 1)
    doc.add_heading("{} Comparison of Feedback received from Self vs Surveyor Type {}".format('.'.join(sec_lst), SurveyorTypeKey2Label[surveyor_type_key]), level=heading_level)
    doc.add_paragraph(f"Following is the comparison summary of feedback received from Self vs Surveyor type {SurveyorTypeKey2Label[surveyor_type_key]} for the Question List Categories")
    doc.add_heading(f"{'.'.join(sec_lst) + '.1'} Comparison Data between Self and Surveyor Type: {SurveyorTypeKey2Label[surveyor_type_key]}", level= heading_level + 1)
    ratings_table = doc.add_table(rows=11, cols = (len(category_names) + 1))
    ratings_table.style = 'Medium Grid 2 Accent 1'
    ratings_table.autofit = False
    ratings_table.allow_autofit = False
    FormatTableBorders(ratings_table)
    hdr_cells = ratings_table.rows[0].cells
    hdr_cells[0].text = 'Surveyor Type'
    for col in range(0, len(category_names)):
      hdr_cells[col + 1].text = category_names[col]
    row_headings = ["Self-Avg", surveyor_type_key + "-Avg", "Self-Mean", surveyor_type_key + "-Mean",
                    "Self-StdDev", surveyor_type_key + "-StdDev", "Self-MaxRating", surveyor_type_key + "-MaxRating",
                    "Self-MinRating", surveyor_type_key + "-MinRating"]
    for row in range(1, len(row_headings) + 1):
      ratings_table.rows[row].cells[0].text = row_headings[row - 1]
    for cat_idx in range(0, len(data_map['DetailedRatings']['Self']['RatingsAveragePerCategory'])):
      ratings_table.rows[1].cells[cat_idx + 1].text = str(data_map['DetailedRatings']['Self']['RatingsAveragePerCategory'][cat_idx])
      ratings_table.rows[2].cells[cat_idx + 1].text = str(data_map['DetailedRatings'][surveyor_type_key]['RatingsAveragePerCategory'][cat_idx])
      ratings_table.rows[3].cells[cat_idx + 1].text = str(data_map['DetailedRatings']['Self']['RatingsMeanPerCategory'][cat_idx])
      ratings_table.rows[4].cells[cat_idx + 1].text = str(data_map['DetailedRatings'][surveyor_type_key]['RatingsMeanPerCategory'][cat_idx])
      ratings_table.rows[5].cells[cat_idx + 1].text = str(data_map['DetailedRatings']['Self']['RatingsStdDeviationPerCategory'][cat_idx])
      ratings_table.rows[6].cells[cat_idx + 1].text = str(data_map['DetailedRatings'][surveyor_type_key]['RatingsStdDeviationPerCategory'][cat_idx])
      ratings_table.rows[7].cells[cat_idx + 1].text = str(data_map['DetailedRatings']['Self']['RatingsMaximum'][cat_idx])
      ratings_table.rows[8].cells[cat_idx + 1].text = str(data_map['DetailedRatings'][surveyor_type_key]['RatingsMaximum'][cat_idx])
      ratings_table.rows[9].cells[cat_idx + 1].text = str(data_map['DetailedRatings']['Self']['RatingsMinimum'][cat_idx])
      ratings_table.rows[10].cells[cat_idx + 1].text = str(data_map['DetailedRatings'][surveyor_type_key]['RatingsMinimum'][cat_idx])
    FormatTableColumnFontProperties(ratings_table.columns[0], 1, 'Verdana', Pt(6), False, False)
    doc.add_heading(f"{'.'.join(sec_lst) + '.2'} Scatter Graph of Ratings Average between Self and Surveyor Type: {SurveyorTypeKey2Label[surveyor_type_key]}", level= heading_level + 1)
    PlotScatterGraphpOfFeedbackAveragesBetweenSelfAndSurveyorType(survey_id, eid, surveyor_type_key, seid_2_info_map, category_names, data_map, doc)
    doc.add_heading(f"{'.'.join(sec_lst) + '.3'} Radar Graph of Ratings Average between Self and Surveyor Type: {SurveyorTypeKey2Label[surveyor_type_key]}", level= heading_level + 1)
    PlotRadarGraphOfFeedbackAveragesBetweenSelfAndSurveyorType(survey_id, eid, surveyor_type_key, seid_2_info_map, category_names, data_map, doc)
    doc.add_heading(f"{'.'.join(sec_lst) + '.4'} Filled Radar Graph of Ratings Average between Self and Surveyor Type: {SurveyorTypeKey2Label[surveyor_type_key]}", level= heading_level + 1)
    PlotFilledRadarGraphpOfFeedbackAveragesBetweenSelfAndSurveyorType(survey_id, eid, surveyor_type_key, seid_2_info_map, category_names, data_map, doc)
    
  
def PlotRadarGraphpOfFeedbackAveragesBySurveyorType(survey_id, eid, seid_2_info_map, category_names, data_map, doc):
  print(f"Creating Feedback graph received for employee:{eid} from Surveyors")
  N = len(category_names)
  theta = radar_factory(N, frame='polygon')
  print(f"Using N={N}, theta={theta}")
  fig, ax = plt.subplots(1, 1, figsize=(8, 8), layout='constrained', subplot_kw=dict(projection='radar'))
  fig.subplots_adjust(wspace=0.25, hspace=0.20, top=0.85, bottom=0.05)
  ax.set_title(f"Radar Graph of Feedback Received from All Surveyors", weight='bold', size='medium', position=(0.5, 1.1),
              horizontalalignment='center', verticalalignment='center')
  employee_name = seid_2_info_map[eid]['Name']    
  legend_strs = []
  for surveyor_type_key in data_map['DetailedRatings']:
    ax.plot(theta, data_map['DetailedRatings'][surveyor_type_key]['RatingsAveragePerCategory'], color=ColorMap[surveyor_type_key])
    legend_strs.append(SurveyorTypeKey2Label[surveyor_type_key])
  ax.set_varlabels(category_names)
  ax.legend(legend_strs, loc=(0.9, .95), labelspacing=0.1, fontsize='small')
  fig.suptitle(f"Survey ID: {survey_id} Average Feedback Received from All Surveyors for employee: {employee_name}")
  image_file = "survey_" + survey_id + employee_name.replace(" ", "_") + "_radar_graph.jpg"
  print(f"Saving image to file:{image_file}")
  plt.savefig(image_file, bbox_inches='tight')
  doc.add_picture(image_file, width=Pt(300))
  if DisplayGraphsFlag:
    plt.show()
  if not KeepTempsFlag:
    os.remove(image_file)


def PlotFilledRadarGraphpOfFeedbackAveragesBySurveyorType(survey_id, eid, seid_2_info_map, category_names, data_map, doc):
  print(f"Creating Feedback graph received for employee:{eid} from Surveyors")
  N = len(category_names)
  theta = radar_factory(N, frame='polygon')
  print(f"Using N={N}, theta={theta}")
  fig, ax = plt.subplots(1, 1, figsize=(8, 8), layout='constrained', subplot_kw=dict(projection='radar'))
  fig.subplots_adjust(wspace=0.25, hspace=0.20, top=0.85, bottom=0.05)
  ax.set_title(f"Radar Graph of Feedback Received from All Surveyors", weight='bold', size='medium', position=(0.5, 1.1),
              horizontalalignment='center', verticalalignment='center')    
  employee_name = seid_2_info_map[eid]['Name']
  legend_strs = []
  for surveyor_type_key in data_map['DetailedRatings']:
    color = ColorMap[surveyor_type_key]
    ax.plot(theta, data_map['DetailedRatings'][surveyor_type_key]['RatingsAveragePerCategory'], color=color)
    ax.fill(theta, data_map['DetailedRatings'][surveyor_type_key]['RatingsAveragePerCategory'], facecolor=color, alpha=0.25, label='_nolegend_')
    legend_strs.append(SurveyorTypeKey2Label[surveyor_type_key])
  ax.set_varlabels(category_names)
  ax.legend(legend_strs, loc=(0.9, .95), labelspacing=0.1, fontsize='small')
  fig.suptitle(f"Survey ID: {survey_id} Average Feedback Received from All Surveyors for employee: {employee_name}")
  image_file = "survey_" + survey_id + employee_name.replace(" ", "_") + "_filled_radar_graph.jpg"
  print(f"Saving image to file:{image_file}")
  plt.savefig(image_file, bbox_inches='tight')
  doc.add_picture(image_file, width=Pt(300))
  if DisplayGraphsFlag:
    plt.show()
  if not KeepTempsFlag:
    os.remove(image_file)
  

def PlotScatterGraphpOFeedbackAveragesBySurveyorType(survey_id, eid, seid_2_info_map, category_names, data_map, doc):  
  print(f"Creating Scatter graph of Average Feedback received for employee:{eid} from Different Surveyor Types")
  print(f"Using category list={category_names}\n")
  fig = plt.figure(figsize=(8.5, 5), layout='constrained')
  ax = fig.subplots()
  print(f"Plotting, Scatter Graph of Average Ratings from Surveyors for CategoryNames=[{category_names}]")
  for surveyor_type_key in data_map['DetailedRatings']:
    print(f"\tPlot surveyor_type={surveyor_type_key}, Averages={data_map['DetailedRatings'][surveyor_type_key]['RatingsAveragePerCategory']}, Type of array={type(data_map['DetailedRatings'][surveyor_type_key]['RatingsAveragePerCategory'])}\n")
    plt.plot(category_names, data_map['DetailedRatings'][surveyor_type_key]['RatingsAveragePerCategory'],
             label = SurveyorTypeKey2Label[surveyor_type_key], color=ColorMap[surveyor_type_key])
  plt.xlabel("Evaluation Category")
  plt.ylabel("Average Feedback Rating")
  employee_name = seid_2_info_map[eid]['Name'] 
  plt.title(f"Survey ID: {survey_id} Average Feedback Received from All Surveyors for employee: {employee_name}")
  plt.legend()
  image_file = "survey_" + survey_id + employee_name.replace(" ", "_") + "_combined_scatter_graph.jpg"
  plt.savefig(image_file, bbox_inches='tight')
  doc.add_picture(image_file, width=Pt(300))
  if DisplayGraphsFlag:
    plt.show()
  if not KeepTempsFlag:
    os.remove(image_file)


def WriteFeedbackAveragesBySurveyorTypeReport(survey_id, eid, seid_2_info_map, category_names, data_map, doc, sec_lst):
  #sec_lst = re.split(r'[.]', sect_str)
  sec_lst[-1] = str(int(sec_lst[-1]) + 1)
  heading_level = len(sec_lst)
  doc.add_heading("{} Table of Rating Averages by Surveyor Type for each Category".format('.'.join(sec_lst)), level=heading_level)
  doc.add_paragraph("Following is the summary of combined feedback received by Surveyor type for the Question List Categories")
  surveyor_types = list(data_map['DetailedRatings'].keys())
  print(f"surveyor_types={surveyor_types}, type={type(surveyor_types)}")
  ratings_table = doc.add_table(rows=(len(surveyor_types) + 1), cols = (len(category_names) + 1))
  ratings_table.style = 'Medium Grid 2 Accent 1'
  ratings_table.autofit = False
  ratings_table.allow_autofit = False
  FormatTableBorders(ratings_table)
  hdr_cells = ratings_table.rows[0].cells
  hdr_cells[0].text = 'Surveyor Type'
  for col in range(0, len(category_names)):
    hdr_cells[col + 1].text = category_names[col]
  row =  1
  for row in range(0, len(surveyor_types)):
    row_cells = ratings_table.rows[row + 1].cells
    row_cells[0].text = SurveyorTypeKey2Label[surveyor_types[row]]
    print(f"Surveyor Type={row_cells[0].text}")
    for cat_idx in range(0, len(data_map['DetailedRatings'][surveyor_types[row]]['RatingsAveragePerCategory'])):
      row_cells[cat_idx + 1].text = str(data_map['DetailedRatings'][surveyor_types[row]]['RatingsAveragePerCategory'][cat_idx])
      print(f"row={row}, cat_idx={cat_idx}, average={row_cells[cat_idx + 1].text}")
  FormatTableColumnFontProperties(ratings_table.columns[0], 1, 'Verdana', Pt(6), False, False)
  doc.add_heading("{} Scatter Graph of Feedback Averages received from each Surveyor Type".format('.'.join(sec_lst)), level=heading_level)
  PlotScatterGraphpOFeedbackAveragesBySurveyorType(survey_id, eid, seid_2_info_map, category_names, data_map, doc)          
  sec_lst[-1] = str(int(sec_lst[-1]) + 1)
  doc.add_heading("{} Radar Graph of Feedback Averages received from each Surveyor Type".format('.'.join(sec_lst)), level=heading_level)
  PlotRadarGraphpOfFeedbackAveragesBySurveyorType(survey_id, eid, seid_2_info_map, category_names, data_map, doc)
  sec_lst[-1] = str(int(sec_lst[-1]) + 1)
  doc.add_heading("{} Filled Radar Graph of Feedback Averages received from each Surveyor Type".format('.'.join(sec_lst)), level=heading_level)
  PlotFilledRadarGraphpOfFeedbackAveragesBySurveyorType(survey_id, eid, seid_2_info_map, category_names, data_map, doc)
        
  
def CreateReport(mysql_con_obj, eid, survey_id):
  print("------------------------------------CreateReport:[start]-----------------------")
  feedback_list = []
  print(f"Obtaining feedback by Surveyor type for employee={eid}, for survey={survey_id}\n")
  GetSurveyFeedbackForEmployee(mysql_con_obj, survey_id, eid, feedback_list)
  seid_2_info_map = {}
  stype_2_info_map = {}
  surveyor_types_list = []
  GetSurveyorInformationForAssignment(mysql_con_obj, survey_id, eid, seid_2_info_map, stype_2_info_map)
  GetOrderedSurveyorList(stype_2_info_map, surveyor_types_list)
  print(f"Using surveyor types in order:{surveyor_types_list}")
  print(f"Obtaining QuestionListId of Survey={survey_id}")
  print("printing surveyor type list:")
  for stype in surveyor_types_list:
    print(f"\tstype=[{stype}")
  question_list_id = GetQuestionListIdOfSurvey(mysql_con_obj, survey_id)
  print(f"Obtaining Question List details for QuestionListId={question_list_id}")
  questionaire_details_map = OrderedDict()
  GetQuestionaireDetailsFromQuestionListId(mysql_con_obj, question_list_id, questionaire_details_map)
  category_list = []
  GetCategoryListFromQuestionaireDetails(questionaire_details_map, category_list)
  data_map = OrderedDict()
  dt = np.dtype([('surveyor_eid', np.unicode_, 16),('type', int), ('category', int), ('question_number', int), ('rating', int)])
  ratings_array = np.array(feedback_list, dtype=dt)
  print("-------------Numpy: Ratings Array-------------- ")
  print(ratings_array)
  ProcessData(ratings_array, data_map, questionaire_details_map)
  print("---------------DataMap-----------------------\n")
  print(data_map)
  surveyee_name = seid_2_info_map[eid]['Name']
  surveyee_dept = seid_2_info_map[eid]['Department']
  surveyee_job_title = seid_2_info_map[eid]['JobTitle']
  document = CreateCustomDocument()
  document.add_heading('Survey360 Report of Employee:' + surveyee_name, 0)
  document.add_page_break()
  section_list = ['1', '0']
  section_str = '.'.join(section_list)
  document.add_heading('{} Introduction'.format(section_str), level=1)
  definition = """A 360 Feedback Survey also called a 360 Survey or a 360 Review is a type of an assessment of the performance of an 
  employee by gathering feedback from multiple  sources such as Self, Reporting Manager(s), Direct Reports, and Peers. It can also be
  sourced from External stakeholders when an employee interacts regularly with external sources such as Suppliers, Vendors, or Contractors."""
  document.add_paragraph(definition)
  section_list[0] = '2'
  section_str = '.'.join(section_list)
  document.add_heading("{} 360 Feedback Survey of employee {}".format(section_str, surveyee_name), level=2)
  p = document.add_paragraph('This document contains the Feedback received on the 360 Survey conducted on Employee ')
  p.add_run(surveyee_name).bold = True
  section_list[-1] = str(int(section_list[-1]) + 1)
  document.add_heading("{} Employee Details".format('.'.join(section_list)), level=2)
  records = (
    ('Employee Name', surveyee_name),
    ('Employee ID', eid),
    ('Department', surveyee_dept),
    ('Job Title', surveyee_job_title)
  )
  employee_info_table = document.add_table(rows=len(records), cols=2)
  employee_info_table.style = 'Medium Grid 2 Accent 1'
  employee_info_table.autofit = False
  employee_info_table.allow_autofit = False
  FormatTableColumnWidth(employee_info_table.columns[0], 914400)
  FormatTableColumnWidth(employee_info_table.columns[0], 1828800)
  FormatTableBorders(employee_info_table)
  for row in range(0, len(records)):
    row_cells = employee_info_table.rows[row].cells
    row_cells[0].text = records[row][0]
    row_cells[1].text = records[row][1]
  FormatTableCellFontProperties(employee_info_table.rows[0].cells[1], "Verdana", Pt(6), False, False)
  section_list[-1] = str(int(section_list[-1]) + 1)  
  document.add_heading("{} Surveyors".format('.'.join(section_list)), level=2)
  document.add_paragraph('The following list of surveyors provided feedback on the survey in addition to self.')      
  #confirm if the names & eids of surveyors should be listed here
  surveyor_info_table = document.add_table(rows=1, cols=5)
  surveyor_info_table.style = 'Medium Grid 2 Accent 1'
  hdr_cells = surveyor_info_table.rows[0].cells
  hdr_cells[0].text = 'Surveyor Name'
  hdr_cells[1].text = 'Surveyor Type'
  hdr_cells[2].text = 'Employee ID'
  hdr_cells[3].text = 'Job Title'
  hdr_cells[4].text = 'Department'
  for stype in stype_2_info_map:
    for seid in stype_2_info_map[stype]:
      row_cells = surveyor_info_table.add_row().cells
      print(f"Adding Surveyor={stype_2_info_map[stype][seid]['Name']}, Type={stype}, EID={seid}\n")
      row_cells[0].text = stype_2_info_map[stype][seid]['Name']
      #row_cells[0].add_paragraph().add_run(text=stype_2_info_map[stype][seid]['Name'], style = None)
      row_cells[1].text = stype
      row_cells[2].text = seid
      row_cells[3].text = stype_2_info_map[stype][seid]['JobTitle']
      row_cells[4].text = stype_2_info_map[stype][seid]['Department']
  FormatTableColumnFontProperties(surveyor_info_table.columns[0], 1, 'Verdana', Pt(6), False, False)
  FormatTableColumnWidth(surveyor_info_table.columns[0], 1200000)
  FormatTableBorders(surveyor_info_table)
  print(f"Number of rows in surveyor_info_table={len(surveyor_info_table.rows)}")
  #FormatTableColumnFontProperties(surveyor_info_table.columns[0], 1, 'Verdana', Pt(6), False, False)
  section_list[-1] = str(int(section_list[-1]) + 1) 
  document.add_heading("{} Evaluation Criteria".format('.'.join(section_list)), level=2)
  section_list.append('1')
  document.add_heading("{} Evaluation Categories".format('.'.join(section_list)), level=3)
  document.add_paragraph().add_run("The employee was evaluated on the following Categories during this survey")
  for category_info_record in category_list:
    document.add_paragraph(category_info_record[1], style='List Number')
  section_list[-1] = str(int(section_list[-1]) + 1)   
  document.add_heading("{} Ratings Scale".format('.'.join(section_list)), level=3)
  rating_records = (
    (1, 'Not demonstrated it at all'),
    (2, '3 out of 10 times the individual has demonstrated it'),
    (3, '4-7 times out of 10 the individual has demonstrated it'),
    (4, '8-9 times out of 10 the individual has demonstrated it'),
    (5, 'Individual has demonstrated in every interaction')
    )
  ratings_info_table = document.add_table(rows=1, cols=2)
  hdr_cells = ratings_info_table.rows[0].cells
  #hdr_cells[0].text = 'Rating'
  hdr_cells[0].add_paragraph().add_run(text='Rating').bold = True
  #hdr_cells[1].text = 'Description'
  hdr_cells[1].add_paragraph().add_run(text='Description').bold = True
  for rating_record in rating_records:
    row_cells = ratings_info_table.add_row().cells
    row_cells[0].text = str(rating_record[0])
    row_cells[1].text = rating_record[1]
  FormatTableColumnWidth(ratings_info_table.columns[0], 1000000)
  FormatTableColumnWidth(ratings_info_table.columns[1], 2200000)
  FormatTableBorders(ratings_info_table)
  section_list[-1] = str(int(section_list[-1]) + 1)
  document.add_heading("{} Surveyor Types".format('.'.join(section_list)), level=3)
  p = document.add_paragraph('The participating employee ')
  p.add_run(surveyee_name).italic = True
  p.add_run(' was assessed by the following categories of surveyor types including the self-evaluation.')
  for surveyor_type in stype_2_info_map:
    document.add_paragraph(surveyor_type, style='List Bullet')
  #document.add_page_break()
  category_names = [e[1].replace(' ', '\n') for e in category_list]  
  section_list = ['3', '0']
  document.add_heading("{} Feedback Ratings Summary".format('.'.join(section_list)), level=1)
  WriteRatingsSummaryReport(survey_id, eid, seid_2_info_map, category_names, ratings_array, data_map, document, section_list)
  section_list = ['4', '0']  
  document.add_heading("{} Detailed Feedback Ratings".format('.'.join(section_list)), level=1)
  document.add_paragraph("""This section contains Feedback Averages received from Surveyor Types relevant to the employee.""")
  WriteFeedbackAveragesBySurveyorTypeReport(survey_id, eid, seid_2_info_map, category_names, data_map, document, section_list)
  section_list = ['5', '0']  
  document.add_heading("{} Detailed Feedback Ratings Comparison between Self and Surveyor Types".format('.'.join(section_list)), level=1)  
  WriteFeedbackComparisonReport(survey_id, eid, seid_2_info_map, category_names, data_map, document, section_list)
  section_list = ['6', '0']
  document.add_heading("{} Detailed Analysis of each Category Type".format('.'.join(section_list)), level=1)  
  WriteCategoryWiseReport(survey_id, eid, seid_2_info_map, ratings_array, surveyor_types_list, category_list,
                          questionaire_details_map, document, section_list)
  document_name = "Survey_" + survey_id + '_' + surveyee_name + '.docx'
  print(f"Saving report in file:{document_name}")
  document.save(document_name)
  print("------------------------------------CreateReport:[end]-----------------------")

  
def Usage():
  print('Usage: reportgen.py [-h|--help] [-d|--display] [-k|--keep] [-c|--conf <appconfig_file>] [-s|--survey <survey_id>] <employee_id>\n')
  print("-h|--help Display Usage and exit.")
  print("-d|--display Display each graph that is created. The application wont proceed until you close the display window.")
  print("""-k|--keep The application generates a large number of temporary image files of the graphs that are created which
 are removed after transfer to the Word document. This option allows you to save these image files.""")
  print("-c|--conf <appconfig_file> For future use. Use the optional application configuration file.")
  print("-s|--survey <survey_id> Mandatory parameter. Used to obtain feedback data for the Survey specified by the id.")
  print('<employee_id> Mandatory parameter, the report is generated for the Employee identified by this id. Note: Employee Id must be in quotes if the Id contains a space, e.g \'os 01\'')
  

if __name__ == '__main__':
  app_config_file = "reportgen.conf"
  survey_id = 0
  eid = None
  try:
    opts, args = getopt.getopt(sys.argv[1:],"hdrs:c:", ["help", "display", "remove", "survey=", "conf="])
  except getopt.GetoptError:
    Usage()
    sys.exit(2)
  for opt, arg in opts:
    if opt in ('-h', '--help'):
      Usage()
      sys.exit(0)
    elif opt in ('-d', '--display'):
      DisplayGraphsFlag = True
      print("Application will display each Graph that is generated")
    elif opt in ('-k', '--keep'):
      KeepTempsFlag = True
      print("Application will not delete the image files created")
    elif opt in ('-c', '--conf'):
      app_config_file = arg
      print("Config file specified:", app_config_file)
    elif opt in ('-s', '--survey'):
      survey_id = arg
  if (len(args) < 1):
    print("Error: Employee Id argument is  mandatory")
    Usage()
    sys.exit(2)
  eid = str(args[0])
  if not survey_id:
    print(f"Invalid surveyId={survey_id}, SurveyId is a mandatory input parameter\n")
    Usage()
    sys.exit(2)
  mysql_con_obj = MySQLConnection(db_user, db_password, db_host, db_name)
  print("Successfully created mysql connection\n")
  
  open_survey_assignments = GetOpenAssignmnetCountForEmployeeInSurvey(mysql_con_obj, survey_id, eid)
  if open_survey_assignments:
    print(f"{open_survey_assignments} Open Survey Assignments exists for Employee={eid} exists in Survey={survey_id}\n")
    sys.exit(2)
  CreateReport(mysql_con_obj, eid, survey_id)
  
    
